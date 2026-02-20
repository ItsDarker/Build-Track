#!/usr/bin/env python3
"""
Convert Markdown to Word Documents (improved version)
Handles formatting: bold, italic, code, links, headings, lists, tables
"""

import os
import re
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_shading(paragraph, color):
    """Add background shading to a paragraph"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    paragraph._element.get_or_add_pPr().append(shading_elm)

def parse_markdown_to_docx(md_content, docx_path):
    """Convert markdown content to Word document"""
    try:
        doc = Document()
        lines = md_content.split('\n')
        i = 0
        in_code_block = False
        code_block_content = []

        while i < len(lines):
            line = lines[i]

            # Handle code blocks
            if line.strip().startswith('```'):
                if not in_code_block:
                    in_code_block = True
                    code_block_content = []
                else:
                    in_code_block = False
                    if code_block_content:
                        p = doc.add_paragraph()
                        add_shading(p, 'E8E8E8')
                        for code_line in code_block_content:
                            run = p.add_run(code_line + '\n')
                            run.font.name = 'Courier New'
                            run.font.size = Pt(10)
                code_block_content = []
                i += 1
                continue

            if in_code_block:
                code_block_content.append(line)
                i += 1
                continue

            # Handle tables
            if line.strip().startswith('|'):
                table_lines = []
                while i < len(lines) and lines[i].strip().startswith('|'):
                    table_lines.append(lines[i])
                    i += 1

                if table_lines:
                    # Parse table
                    rows = [re.split(r'\|', row.strip('|').strip()) for row in table_lines if row.strip()]
                    if len(rows) >= 2:
                        # Skip separator row
                        is_sep = all(':' in cell or '-' in cell for cell in rows[1][0].split())
                        if is_sep and len(rows) > 2:
                            rows = [rows[0]] + rows[2:]

                        table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                        table.style = 'Light Grid Accent 1'

                        # Format header row
                        header_cells = table.rows[0].cells
                        for cell in header_cells:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.bold = True

                        # Fill table data
                        for row_idx, row in enumerate(rows):
                            cells = table.rows[row_idx].cells
                            for col_idx, cell_content in enumerate(row):
                                if col_idx < len(cells):
                                    cells[col_idx].text = cell_content.strip()
                continue

            # Handle headings
            if line.startswith('#'):
                level = len(line) - len(line.lstrip('#'))
                text = line.lstrip('#').strip()
                if text:
                    p = doc.add_heading(text, level=min(level, 6))
                i += 1
                continue

            # Handle blank lines
            if not line.strip():
                doc.add_paragraph()
                i += 1
                continue

            # Handle bullet lists
            if line.strip().startswith('-') and not line.startswith('---'):
                list_items = []
                while i < len(lines) and lines[i].strip().startswith('-') and not lines[i].startswith('---'):
                    item = lines[i].strip().lstrip('-').strip()
                    list_items.append(item)
                    i += 1

                for item in list_items:
                    doc.add_paragraph(item, style='List Bullet')
                continue

            # Handle numbered lists
            if re.match(r'^\d+\.', line.strip()):
                list_items = []
                while i < len(lines) and re.match(r'^\d+\.', lines[i].strip()):
                    item = re.sub(r'^\d+\.\s*', '', lines[i].strip())
                    list_items.append(item)
                    i += 1

                for item in list_items:
                    doc.add_paragraph(item, style='List Number')
                continue

            # Handle regular paragraphs with inline formatting
            p = doc.add_paragraph()
            text = line.strip()

            if text:
                # Parse inline formatting
                # This regex-based approach handles bold, italic, code, links
                pattern = r'(\*\*.*?\*\*|__.*?__|`.*?`|\[.*?\]\(.*?\)|_.*?_|\*.*?\*)'
                parts = re.split(pattern, text)

                for part in parts:
                    if not part:
                        continue

                    # Bold
                    if part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part[2:-2])
                        run.font.bold = True
                    elif part.startswith('__') and part.endswith('__'):
                        run = p.add_run(part[2:-2])
                        run.font.bold = True

                    # Italic
                    elif (part.startswith('_') and part.endswith('_') and len(part) > 2):
                        run = p.add_run(part[1:-1])
                        run.font.italic = True
                    elif (part.startswith('*') and part.endswith('*') and len(part) > 2 and not part.startswith('**')):
                        run = p.add_run(part[1:-1])
                        run.font.italic = True

                    # Inline code
                    elif part.startswith('`') and part.endswith('`'):
                        run = p.add_run(part[1:-1])
                        run.font.name = 'Courier New'
                        run.font.color.rgb = RGBColor(192, 0, 0)

                    # Links
                    elif re.match(r'\[.*?\]\(.*?\)', part):
                        match = re.match(r'\[(.*?)\]\((.*?)\)', part)
                        if match:
                            link_text = match.group(1)
                            link_url = match.group(2)
                            run = p.add_run(link_text)
                            run.font.color.rgb = RGBColor(0, 0, 255)
                            run.font.underline = True

                    # Plain text
                    else:
                        p.add_run(part)

            i += 1

        doc.save(docx_path)
        return True
    except Exception as e:
        print(f"Error converting: {str(e)}", file=sys.stderr)
        return False

def convert_all_markdown_files():
    """Convert all markdown files in Docs folder to Word"""
    docs_dir = Path("C:/Personal Files/BuildTrack/Docs")
    output_dir = docs_dir / "Word Documents"

    # Create output directory
    output_dir.mkdir(exist_ok=True)

    # Find all markdown files
    md_files = sorted(docs_dir.glob("*.md"))

    print(f"Found {len(md_files)} markdown files")
    print(f"Output directory: {output_dir}\n")

    success_count = 0
    failed_files = []

    for md_file in md_files:
        print(f"Converting: {md_file.name}...", end=" ", flush=True)

        try:
            # Read markdown
            with open(md_file, 'r', encoding='utf-8') as f:
                content = f.read()

            # Convert to docx
            output_file = output_dir / md_file.with_suffix('.docx').name
            if parse_markdown_to_docx(content, str(output_file)):
                print("[OK]")
                success_count += 1
            else:
                print("[FAILED]")
                failed_files.append(md_file.name)

        except Exception as e:
            print(f"[ERROR] {str(e)}")
            failed_files.append(md_file.name)

    print("\n" + "="*60)
    print("Conversion Summary:")
    print(f"[OK] Successful: {success_count}/{len(md_files)}")
    print(f"[FAILED] Failed: {len(failed_files)}/{len(md_files)}")

    if failed_files:
        print("\nFailed files:")
        for f in failed_files:
            print(f"  - {f}")

    print(f"\nWord documents saved to: {output_dir}")

if __name__ == "__main__":
    convert_all_markdown_files()
